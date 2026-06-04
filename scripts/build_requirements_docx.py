from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "requirements-definition.md"
OUTPUT = ROOT / "docs" / "경남차유리_ERP_요구사항_정의서_v0.1.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
BLACK = RGBColor(25, 31, 40)
BORDER = "D9E2F3"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F8FAFC"
WHITE = "FFFFFF"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: float, color: RGBColor = BLACK, bold: bool = False):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_inches: list[float], indent_dxa: int = 120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = sum(int(w * 1440) for w in widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_inches[min(idx, len(widths_inches) - 1)]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_field(paragraph, field_name: str):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_formatted_text(paragraph, text: str, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    text = text.replace("`", "")
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run_bold = bold
        clean = part
        if part.startswith("**") and part.endswith("**"):
            clean = part[2:-2]
            run_bold = True
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, color=color, bold=run_bold)


def add_para(doc: Document, text: str = "", style: str | None = None, before: float = 0, after: float = 6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        add_formatted_text(p, text)
    return p


def add_callout(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5], indent_dxa=0)
    set_table_borders(table, color=BORDER, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_formatted_text(p, text, color=DARK_BLUE, bold=True)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=10.5, color=DARK_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table_widths(header: list[str]) -> list[float]:
    n = len(header)
    joined = " ".join(header)
    if n == 2:
        return [1.55, 4.95]
    if n == 3:
        if "모델" in joined or "목적" in joined:
            return [1.6, 3.65, 1.25]
        return [1.5, 3.9, 1.1]
    if n == 4:
        if header[0] == "ID":
            return [1.05, 3.75, 0.8, 0.9]
        return [1.45, 2.55, 1.05, 1.45]
    if n == 5:
        return [1.0, 2.55, 0.8, 1.0, 1.15]
    return [6.5 / n for _ in range(n)]


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, table_widths(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            else:
                set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, cols - 2, cols - 1) or len(value) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_text(p, value, size=8.8 if cols >= 4 else 9.2, bold=(r_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip()
        parts = [part.strip() for part in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def collect_headings(lines: list[str]) -> list[str]:
    headings = []
    for line in lines:
        if line.startswith("## ") and not line.startswith("### "):
            headings.append(line[3:].strip())
    return headings


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5, BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    set_style_font(styles["Heading 1"], 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        set_style_font(styles[list_style], 10.5, BLACK)
        styles[list_style].paragraph_format.space_after = Pt(4)
        styles[list_style].paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_text(header, "경남차유리 ERP 요구사항 정의서", size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(footer, "Page ", size=8.5, color=MUTED)
    add_field(footer, "PAGE")


def add_cover(doc: Document, title: str, headings: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    add_formatted_text(p, "요구사항 정의서", size=28, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    add_formatted_text(p, "경남차유리 업무관리 ERP", size=15, color=MUTED, bold=True)

    meta = doc.add_table(rows=5, cols=2)
    set_table_width(meta, [1.4, 5.1], indent_dxa=0)
    set_table_borders(meta, color=BORDER)
    metadata = [
        ("문서 버전", "v0.1 초안"),
        ("작성일", "2026-05-28"),
        ("대상 프로젝트", "경남차유리 업무관리 ERP"),
        ("문서 목적", "초기 요구사항, 프로토타입 설계, Notion 피드백을 통합한 MVP 개발 기준 정리"),
        ("작성 기준", "단일 매장 내부 사용, 빠른 입력, 견적-작업-정산 연결"),
    ]
    for idx, (label, value) in enumerate(metadata):
        meta.cell(idx, 0).text = label
        meta.cell(idx, 1).text = value
        for cell in meta.row_cells(idx):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if cell == meta.cell(idx, 0):
                set_cell_shading(cell, HEADER_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=(cell == meta.cell(idx, 0)))

    add_para(doc, "", after=8)
    add_callout(
        doc,
        "본 문서는 개발 착수 전 범위와 우선순위를 맞추기 위한 초안입니다. "
        "외부 시스템 계정, 비밀번호, 개인 연락처 등 민감 정보는 의도적으로 제외했습니다.",
    )

    doc.add_page_break()
    toc = doc.add_paragraph(style="Heading 1")
    add_formatted_text(toc, "목차")
    for heading in headings:
        p = doc.add_paragraph(style="List Number")
        add_formatted_text(p, heading)
    doc.add_page_break()


def build():
    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = raw_lines[0].lstrip("# ").strip()
    headings = collect_headings(raw_lines)

    doc = Document()
    configure_document(doc)
    add_cover(doc, title, headings)

    start = 0
    for idx, line in enumerate(raw_lines):
        if line.startswith("## "):
            start = idx
            break

    lines = raw_lines[start:]
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            table_rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, table_rows)
            continue

        if stripped.startswith("> "):
            add_callout(doc, stripped[2:].strip())
            idx += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_formatted_text(p, stripped[4:].strip())
            idx += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_formatted_text(p, stripped[3:].strip())
            idx += 1
            continue

        if stripped.startswith("# "):
            idx += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:].strip())
            idx += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, re.sub(r"^\d+\.\s+", "", stripped))
            idx += 1
            continue

        p = add_para(doc, stripped)
        idx += 1

    # Final blank paragraph keeps the last table from feeling pinned to the page edge.
    add_para(doc, "")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
